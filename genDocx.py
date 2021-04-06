# -*- coding: utf-8 -*-
'''
Created on 2019年6月12日

@author: BugTan
'''


from docx import Document

document = Document()
document.add_heading('大傻子ing of the universe')
document.add_heading('The role of dolphins', level=2)
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

paragraph = document.add_paragraph('dadasdasd')

document.save('test.docx')
